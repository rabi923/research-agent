import os
from fpdf import FPDF
from docx import Document
from openpyxl import Workbook
from typing import List, Dict
from research_agent.core.llm.base import BaseLLMProvider
from research_agent.core.llm.factory import get_llm_provider
from research_agent.config.settings import settings

class AnalyzerAgent:
    def __init__(self):
        self.llm: BaseLLMProvider = get_llm_provider()
        self.output_dir = settings.REPORT_OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    def analyze(self, query: str, collected_data: Dict, history: List[Dict] = [], requested_formats: List[str] = ["pdf"], status_callback=None) -> Dict:
        """
        Analyze collected data and generate a report.
        """
        if status_callback:
            status_callback(f"Analyzer Agent: Analyzing collected data for '{query}'...")
        else:
            print(f"Analyzer Agent: Analyzing collected data for '{query}'...")
        
        context = collected_data.get("context", "")
        sources = collected_data.get("sources", [])
        documents = collected_data.get("documents", [])

        # Step 1: Generate Report Content (Markdown)
        report_prompt = f"""
        You are an expert Analyst. Your goal is to synthesize the collected information into a detailed, real-time report.
        
        User Query: {query}
        
        Collected Information:
        {context}
        
        Please generate a comprehensive report in Markdown format.
        Structure the report with:
        - Executive Summary
        - Detailed Analysis (broken down by key topics)
        - Key Findings
        - Conclusion
        
        Do not hallucinate. Base your report strictly on the collected information.
        Do NOT include a references or documents section - this will be added separately.
        """
        
        if status_callback:
            status_callback("Generating comprehensive report...")
        print("DEBUG: Calling LLM for report generation...")
        try:
            report_content = self.llm.generate(report_prompt, history=history, system_prompt="You are a helpful analyst. Write detailed reports.")
            print("DEBUG: Report generated successfully.")
        except Exception as e:
            print(f"DEBUG: Report Generation Failed: {e}")
            report_content = f"Error generating report: {e}\n\nPlease check your API keys and try again."
        
        # Append Sources to Report Body (properly formatted)
        if sources:
            report_content += "\n\n## References\n\n"
            for source in sources:
                report_content += f"- [{source['title']}]({source['url']})\n"

        # Step 2: Generate Files
        if status_callback:
            status_callback("Creating PDF report...")
        pdf_path = self._generate_pdf(query, report_content)
        
        docx_path = None
        if "docx" in requested_formats:
            if status_callback:
                status_callback("Creating DOCX report...")
            docx_path = self._generate_docx(query, report_content)
            
        excel_path = None
        if "excel" in requested_formats:
            if status_callback:
                status_callback("Creating Excel data sheet...")
            excel_path = self._generate_excel(query, sources, documents)
        
        return {
            "report_content": report_content,
            "pdf_path": pdf_path,
            "docx_path": docx_path,
            "excel_path": excel_path,
            "sources": sources,
            "documents": documents
        }

    def _generate_pdf(self, title: str, content: str) -> str:
        """
        Convert Markdown content to PDF with improved formatting.
        """
        import re
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, f"Research Report: {title}", 0, 1, 'C')
        pdf.ln(10)
        
        # Content
        pdf.set_font("Arial", size=11)
        in_references = False
        
        for line in content.split('\n'):
            safe_line = line.encode('latin-1', 'replace').decode('latin-1')
            
            # Check if we're entering References section
            if line.startswith('## References'):
                in_references = True
                pdf.ln(5)
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, 'References', 0, 1)
                pdf.ln(3)
                pdf.set_font("Arial", size=11)
                continue
            
            # Handle section headers
            if line.startswith('# '):
                pdf.ln(5)
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, safe_line.replace('# ', ''), 0, 1)
                pdf.set_font("Arial", size=11)
            elif line.startswith('## '):
                pdf.ln(3)
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, safe_line.replace('## ', ''), 0, 1)
                pdf.set_font("Arial", size=11)
            elif line.startswith('- '):
                # Handle bullet points (especially in References)
                if in_references:
                    # Extract link from Markdown: [title](url)
                    match = re.search(r'\[(.*?)\]\((.*?)\)', line)
                    if match:
                        link_text = match.group(1)
                        link_url = match.group(2)
                        safe_text = link_text.encode('latin-1', 'replace').decode('latin-1')
                        
                        # Create clickable bullet with link
                        pdf.set_x(15)
                        pdf.set_font("Arial", 'U', 11)  # Underline for links
                        pdf.set_text_color(0, 0, 255)   # Blue color for links
                        
                        # Add bullet
                        pdf.write(6, chr(149) + " ")
                        
                        # Add clickable link
                        pdf.cell(0, 6, safe_text, 0, 1, link=link_url)
                        
                        # Reset formatting
                        pdf.set_font("Arial", size=11)
                        pdf.set_text_color(0, 0, 0)
                        pdf.ln(2)  # Add spacing between references
                    else:
                        pdf.set_x(15)
                        pdf.multi_cell(0, 6, chr(149) + " " + safe_line.replace('- ', ''))
                else:
                    # Regular bullet points
                    pdf.set_x(15)
                    pdf.multi_cell(0, 6, chr(149) + " " + safe_line.replace('- ', ''))
            else:
                if line.strip():  # Only add non-empty lines
                    pdf.multi_cell(0, 6, safe_line)
                
        filename = f"{self._sanitize_filename(title)}_report.pdf"
        filepath = os.path.join(self.output_dir, filename)
        pdf.output(filepath)
        return filepath

    def _generate_docx(self, title: str, content: str) -> str:
        """
        Generate DOCX report.
        """
        doc = Document()
        doc.add_heading(f"Research Report: {title}", 0)
        
        for line in content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
            else:
                if line.strip():
                    doc.add_paragraph(line)
                    
        filename = f"{self._sanitize_filename(title)}_report.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def _generate_excel(self, title: str, sources: List[Dict], documents: List[Dict]) -> str:
        """
        Generate Excel sheet with sources and documents.
        """
        wb = Workbook()
        
        # Sources Sheet
        ws_sources = wb.active
        ws_sources.title = "Sources"
        ws_sources.append(["Title", "URL"])
        for source in sources:
            ws_sources.append([source.get("title"), source.get("url")])
            
        # Documents Sheet
        ws_docs = wb.create_sheet("Documents")
        ws_docs.append(["Title", "Type", "URL"])
        for doc in documents:
            ws_docs.append([doc.get("title"), doc.get("type"), doc.get("url")])
            
        filename = f"{self._sanitize_filename(title)}_data.xlsx"
        filepath = os.path.join(self.output_dir, filename)
        wb.save(filepath)
        return filepath

    def _sanitize_filename(self, filename: str) -> str:
        return "".join([c for c in filename if c.isalpha() or c.isdigit() or c=='_' or c=='.']).rstrip()
